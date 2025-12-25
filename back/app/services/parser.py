import pdfplumber
from docx import Document
from io import BytesIO
from typing import Dict, Any


def parse_pdf(file_content: bytes) -> str:
    """Parse PDF file and extract text content"""
    text_content = []

    with pdfplumber.open(BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)

    return "\n\n".join(text_content)


def parse_docx(file_content: bytes) -> str:
    """Parse Word document and extract text content"""
    doc = Document(BytesIO(file_content))
    text_content = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))

    return "\n\n".join(text_content)
